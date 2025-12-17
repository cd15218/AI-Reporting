from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet

def export_pdf(text, charts):
    pdf = SimpleDocTemplate("report.pdf")
    styles = getSampleStyleSheet()
    content = [Paragraph(text.replace("\n", "<br/>"), styles["Normal"])]
    pdf.build(content)

from docx import Document

def export_docx(text, charts):
    doc = Document()
    doc.add_heading("AI Generated Report", level=1)
    doc.add_paragraph(text)
    doc.save("report.docx")
